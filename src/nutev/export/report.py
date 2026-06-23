"""Dissertation/article export: turn a finished run into ready-to-use material.

Reads the pipeline outputs (knowledge base, claims, summaries) and writes a
well-organized ``09_report/`` bundle: an index (INDICE.md) that maps every file
to where it goes in a dissertation, a single multi-sheet workbook
(NUTEV_DISSERTACAO.xlsx), a chapter-structured report (Markdown + optional Word),
an executive summary, and references in BibTeX + RIS.

Nothing here invents findings: it extracts, counts and organizes what the run
produced. Claims/recommendations are surfaced with the human-review caveat.
"""
from __future__ import annotations

import ast
import json
import re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd


def _read_jsonl(path: Path) -> list[dict]:
    if not path.exists():
        return []
    out = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            out.append(json.loads(line))
        except Exception:
            continue
    return out


def _read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_csv(path)
    except Exception:
        return pd.DataFrame()


def _read_json(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _as_list(value) -> list:
    if isinstance(value, list):
        return value
    if isinstance(value, str) and value.strip().startswith("["):
        try:
            v = ast.literal_eval(value)
            return v if isinstance(v, list) else []
        except Exception:
            return []
    return [value] if value not in (None, "", "nan") else []


def _clean(value) -> str:
    """Collapse whitespace/newlines so a value fits on one BibTeX/RIS line."""
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _year(value) -> str:
    s = _clean(value)
    if s.endswith(".0"):  # pandas reads a year column with blanks as float64
        s = s[:-2]
    return s if s.isdigit() else ""


_LATEX = {"&": r"\&", "%": r"\%", "#": r"\#", "_": r"\_", "$": r"\$", "~": " ", "^": " "}


def _bib_escape(value) -> str:
    s = _clean(value).replace("{", "(").replace("}", ")")
    return "".join(_LATEX.get(ch, ch) for ch in s)


def _first_author_surname(authors: str) -> str:
    if not authors or str(authors).strip() in ("", "nan", "None"):
        return "Anon"
    first = re.split(r"\s*;\s*", str(authors).strip())[0]
    if "," in first:  # "Silva, J" -> Silva
        surname = first.split(",")[0]
    else:
        parts = first.split()
        if parts and len(parts[-1]) <= 2 and parts[-1].isupper():
            surname = parts[0]  # "Silva J" (surname + initials)
        else:
            surname = parts[-1] if parts else "Anon"  # "John Silva"
    return re.sub(r"[^A-Za-z]", "", surname) or "Anon"


def _citation_key(rec: dict, used: set[str]) -> str:
    base = re.sub(r"[^A-Za-z0-9]", "", f"{_first_author_surname(rec.get('authors', ''))}{_year(rec.get('year')) or 'sd'}") or "ref"
    candidate, n = base, 1
    while candidate in used:
        n += 1
        candidate = f"{base}{n}"  # always-valid suffix: base, base2, base3, ...
    used.add(candidate)
    return candidate


def _bibtex(records: list[dict]) -> str:
    blocks = []
    for rec in records:
        key = rec.get("_key") or "ref"
        fields = {
            "title": _bib_escape(rec.get("title", "")),
            "author": _bib_escape(rec.get("authors", "")) or "Autor desconhecido",
            "year": _year(rec.get("year")),
            "journal": _bib_escape(rec.get("journal", "")),
            "doi": _clean(rec.get("doi", "")),
            "url": _clean(rec.get("url", "")),
        }
        lines = [f"@article{{{key},"]
        lines += [f"  {k} = {{{v}}}," for k, v in fields.items() if v]
        lines.append("}")
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks) + ("\n" if blocks else "")


def _ris(records: list[dict]) -> str:
    blocks = []
    for rec in records:
        lines = ["TY  - JOUR"]
        if rec.get("_key"):
            lines.append(f"ID  - {rec['_key']}")
        if _clean(rec.get("title")):
            lines.append(f"TI  - {_clean(rec.get('title'))}")
        for author in re.split(r"\s*;\s*", str(rec.get("authors", "") or "")):
            if _clean(author):
                lines.append(f"AU  - {_clean(author)}")
        if _year(rec.get("year")):
            lines.append(f"PY  - {_year(rec.get('year'))}")
        if _clean(rec.get("journal")):
            lines.append(f"JO  - {_clean(rec.get('journal'))}")
        if _clean(rec.get("doi")):
            lines.append(f"DO  - {_clean(rec.get('doi'))}")
        if _clean(rec.get("url")):
            lines.append(f"UR  - {_clean(rec.get('url'))}")
        lines.append("ER  - ")
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks) + ("\n" if blocks else "")


def _top(counter: Counter, n: int = 10) -> list[tuple[str, int]]:
    return [(k, v) for k, v in counter.most_common(n) if k not in ("", "nan", "None", None)]


def _count_df(counter: Counter, col: str) -> pd.DataFrame:
    return pd.DataFrame([{col: k, "n_estudos": v} for k, v in _top(counter, 100)])


def _bullets(pairs: list[tuple[str, int]]) -> str:
    return "\n".join(f"- {k}: {v}" for k, v in pairs) or "- (sem dados)"


def _write_docx(path: Path, *, title: str, caveat: str, sections: list[tuple[str, str]], studies: pd.DataFrame) -> bool:
    """Best-effort Word export; silently skipped if python-docx isn't installed."""
    try:
        import docx
    except Exception:
        return False
    doc = docx.Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(caveat)
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    doc.add_heading("Estudos incluídos", level=1)
    cols = ["titulo", "ano", "periodico", "pais", "nivel_evidencia", "doi"]
    cols = [c for c in cols if c in studies.columns]
    table = doc.add_table(rows=1, cols=len(cols))
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = c
    for _, row in studies.head(200).iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(row.get(c, ""))[:200]
    doc.save(str(path))
    return True


def build_dissertation_report(project_root: Path) -> dict:
    project_root = Path(project_root)
    corpus = _read_jsonl(project_root / "11_knowledge_base" / "corpus.jsonl")
    if not corpus:
        md = _read_csv(project_root / "02_metadata" / "metadata_master.csv")
        corpus = md.fillna("").to_dict("records") if not md.empty else []
    if not corpus:
        return {"available": False, "message": "Nenhum documento encontrado. Rode o pipeline primeiro (▶ Rodar pipeline)."}

    summary = _read_json(project_root / "07_logs" / "run_summary.json")
    claims = _read_csv(project_root / "06_tables" / "NUTEV_EVIDENCE_CLAIMS.csv")

    # Stable citation key per document, reused across .bib/.ris, the studies table
    # and the findings table — so every row maps cleanly to one reference.
    used_keys: set[str] = set()
    key_by_doc: dict[str, dict] = {}
    for rec in corpus:
        rec["_key"] = _citation_key(rec, used_keys)
        key_by_doc[str(rec.get("document_id", ""))] = {
            "key": rec["_key"], "journal": rec.get("journal", ""), "doi": rec.get("doi", ""),
            "url": rec.get("url", ""), "year": rec.get("year", ""), "title": rec.get("title", ""),
        }

    out_dir = project_root / "09_report"
    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[str] = []

    def _write(name: str, text: str) -> None:
        (out_dir / name).write_text(text, encoding="utf-8")
        written.append(f"09_report/{name}")

    # --- references ---
    _write("referencias.bib", _bibtex(corpus))
    _write("referencias.ris", _ris(corpus))

    # --- per-article table (included studies + evaluation) ---
    studies = pd.DataFrame([
        {
            "chave_citacao": r.get("_key", ""),
            "titulo": r.get("title", ""), "autores": r.get("authors", ""), "ano": _year(r.get("year")),
            "periodico": r.get("journal", ""), "pais": "; ".join(str(x) for x in _as_list(r.get("countries"))),
            "tema": "; ".join(str(x) for x in _as_list(r.get("domains"))),
            "tipo_evidencia": r.get("evidence_type", ""), "nivel_evidencia": r.get("evidence_tier", ""),
            "qualidade_periodico": r.get("journal_quality_score", ""), "citacoes": r.get("cited_by_count", ""),
            "relevancia": r.get("relevance_score", ""), "workstream": r.get("workstream", ""),
            "doi": r.get("doi", ""), "url": r.get("url", ""),
        }
        for r in corpus
    ])
    if "tema" in studies.columns:
        studies = studies.sort_values(["tema", "ano"], na_position="last").reset_index(drop=True)
    studies.to_csv(out_dir / "estudos_incluidos.csv", index=False)
    written.append("09_report/estudos_incluidos.csv")

    # --- aggregates ---
    tiers, etypes, countries, domains, workstreams, journals = (Counter() for _ in range(6))
    for r in corpus:
        tiers[str(r.get("evidence_tier", "") or "n/d")] += 1
        etypes[str(r.get("evidence_type", "") or "n/d")] += 1
        workstreams[str(r.get("workstream", "") or "n/d")] += 1
        if r.get("journal"):
            journals[str(r.get("journal"))] += 1
        for c in _as_list(r.get("countries")):
            countries[str(c)] += 1
        for d in _as_list(r.get("domains")):
            domains[str(d)] += 1
    years = sorted({int(_year(r.get("year"))) for r in corpus if _year(r.get("year"))})
    year_counts = Counter(_year(r.get("year")) for r in corpus if _year(r.get("year")))
    most_cited = sorted(
        [r for r in corpus if str(r.get("cited_by_count", "")).strip().isdigit()],
        key=lambda r: int(r["cited_by_count"]), reverse=True,
    )[:10]

    n_docs = summary.get("records", len(corpus))
    n_unique = summary.get("curated_unique_documents", len(corpus))
    claims_total = summary.get("evidence_claims_total", len(claims))
    claims_support = summary.get("evidence_claims_supported", 0)
    claims_review = summary.get("evidence_claims_needs_review", 0)
    recs_total = summary.get("recommendation_candidates_total", 0)
    conflicts = summary.get("conflicting_evidence_total", 0)
    period = f"{years[0]}–{years[-1]}" if years else "n/d"
    gen_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    supported = claims[claims.get("claim_status").astype(str).str.contains("support", case=False, na=False)] if "claim_status" in claims.columns else pd.DataFrame()

    # --- principais achados: cada claim apoiado ligado ao artigo + chave de citação ---
    achados = pd.DataFrame()
    if not supported.empty:
        ach_rows = []
        for _, c in supported.iterrows():
            doc = key_by_doc.get(str(c.get("document_id", "")), {})
            ach_rows.append({
                "tema": "; ".join(str(x) for x in (_as_list(c.get("nutev_domains")) or _as_list(c.get("evidence_lenses")))),
                "achado": _clean(c.get("claim_text", "")),
                "citacao_textual": _clean(c.get("exact_quote", "")),
                "artigo": _clean(c.get("title", "") or doc.get("title", "")),
                "ano": _year(c.get("year")) or _year(doc.get("year")),
                "periodico": doc.get("journal", ""),
                "chave_citacao": doc.get("key", ""),
                "doi": doc.get("doi", ""),
                "url": c.get("source_url", "") or doc.get("url", ""),
            })
        achados = pd.DataFrame(ach_rows).sort_values(["tema", "ano"], na_position="last").reset_index(drop=True)
        achados.to_csv(out_dir / "achados_principais.csv", index=False)
        written.append("09_report/achados_principais.csv")

    # --- single organized workbook (one file, several sheets) ---
    resumo_df = pd.DataFrame([
        ("Documentos", n_docs), ("Únicos curados", n_unique), ("Claims (total)", claims_total),
        ("Claims apoiados", claims_support), ("Claims p/ revisão", claims_review),
        ("Recomendações", recs_total), ("Conflitos", conflicts), ("Período", period),
        ("Países distintos", len(countries)), ("Gerado em", gen_at),
    ], columns=["metrica", "valor"])
    try:
        with pd.ExcelWriter(out_dir / "NUTEV_DISSERTACAO.xlsx", engine="openpyxl") as xl:
            resumo_df.to_excel(xl, sheet_name="Resumo", index=False)
            studies.to_excel(xl, sheet_name="Estudos_incluidos", index=False)
            _count_df(domains, "tema").to_excel(xl, sheet_name="Por_tema", index=False)
            _count_df(countries, "pais").to_excel(xl, sheet_name="Por_pais", index=False)
            _count_df(year_counts, "ano").to_excel(xl, sheet_name="Por_ano", index=False)
            _count_df(journals, "periodico").to_excel(xl, sheet_name="Por_periodico", index=False)
            _count_df(workstreams, "workstream").to_excel(xl, sheet_name="Por_workstream", index=False)
            if not achados.empty:
                achados.to_excel(xl, sheet_name="Principais_achados", index=False)
        written.append("09_report/NUTEV_DISSERTACAO.xlsx")
    except Exception:
        pass

    # --- results organized by theme (maps to dissertation result sub-sections) ---
    by_theme_md = []
    for dom, n in _top(domains, 12):
        sample = [r for r in corpus if dom in [str(x) for x in _as_list(r.get("domains"))]][:3]
        items = "\n".join(f"  - {s.get('title','(sem título)')} ({s.get('year','s.d.')})" for s in sample)
        by_theme_md.append(f"**{dom}** — {n} estudos\n{items}")
    by_theme_section = "\n\n".join(by_theme_md) or "(sem temas)"

    top_cited_md = "\n".join(
        f"{i}. {r.get('title','(sem título)')} ({r.get('year','s.d.')}) — {r.get('cited_by_count','0')} citações"
        for i, r in enumerate(most_cited, 1)
    ) or "- (sem dados de citação)"

    achados_md = []
    if not achados.empty:
        for tema, grp in list(achados.groupby("tema"))[:10]:
            items = "\n".join(
                f"  - {str(row['achado'])[:200]} — _{str(row['artigo'])[:60]} ({row['ano']})_ **[{row['chave_citacao']}]**"
                for _, row in grp.head(3).iterrows()
            )
            achados_md.append(f"**{tema or 'geral'}**\n{items}")
    achados_section = "\n\n".join(achados_md) or "(veja `achados_principais.csv` — todos os achados com artigo e chave de citação)"

    methods_txt = (
        f"Busca multilíngue e multibase (workstreams: {', '.join(summary.get('workstreams', []) or ['n/d'])}). "
        f"Provedores: {summary.get('providers_started', 'n/d')} execuções de consulta. "
        f"Downloads de texto completo: {summary.get('downloads_ok', 0)} ok / {summary.get('downloads_failed', 0)} apenas metadados. "
        "Fluxo PRISMA, estratégia de busca e protocolo (PROSPERO) em 08_docs/ e 06_tables/NUTEV_PRISMA2020_FLOW.csv."
    )
    results_txt = (
        f"Documentos: {n_docs} (únicos: {n_unique}). Claims: {claims_total} "
        f"({claims_support} apoiados, {claims_review} p/ revisão). Recomendações: {recs_total}. "
        f"Conflitos: {conflicts}. Período: {period}."
    )
    eval_txt = (
        f"{len(corpus)} registros únicos, {len(countries)} países, {len(years)} anos. "
        "Distribuição por nível de evidência e qualidade de periódico nas abas da planilha; "
        "cruze com config/predatory_journals.json antes de incluir."
    )

    # --- executive summary ---
    _write("RESUMO_EXECUTIVO.md", f"""# Resumo executivo — NutEV

_Gerado em {gen_at}. **Material de apoio: nenhuma recomendação é final sem revisão humana.**_

- Documentos: **{n_docs}** ({n_unique} únicos) · Período **{period}**
- Claims: **{claims_total}** ({claims_support} apoiados · {claims_review} p/ revisão)
- Recomendações: **{recs_total}** · Conflitos: **{conflicts}**
- Países (top): {', '.join(f'{k} ({v})' for k, v in _top(countries, 5)) or 'n/d'}
- Temas (top): {', '.join(f'{k} ({v})' for k, v in _top(domains, 5)) or 'n/d'}

Comece pelo **INDICE.md**. Tudo está em `09_report/`.
""")

    # --- consolidated, chapter-structured report ---
    _write("RELATORIO_DISSERTACAO.md", f"""# Revisão de evidências NutEV — relatório para dissertação/artigos

_Gerado automaticamente em {gen_at}._
> ⚠️ Material de **apoio** da coleta automatizada. Claims e recomendações são
> **candidatos** e exigem **validação humana** antes do uso final.

## 1. Métodos (resumo)
{methods_txt}

## 2. Resultados quantitativos
{results_txt}

### Por nível de evidência
{_bullets(_top(tiers))}

### Por país (top 10)
{_bullets(_top(countries))}

### Por ano
{_bullets(_top(year_counts))}

## 3. Avaliação dos artigos
{eval_txt}

### Mais citados (proxy de impacto)
{top_cited_md}

## 4. Resultados por tema
{by_theme_section}

## 5. Principais achados (claims apoiados)
**Todos** os achados apoiados estão em **`achados_principais.csv`** (aba `Principais_achados`) —
cada um já traz o artigo, o ano e a **chave de citação** (escreva o achado e cite pela chave/DOI).
Amostra por tema:

{achados_section}

## 6. Recomendações candidatas
{recs_total} candidata(s) em `06_tables/NUTEV_RECOMENDACOES_OPERACIONAIS.xlsx`. **Não use sem revisão humana.**

## 7. Lacunas e conflitos
Conflitos: {conflicts} (`06_tables/NUTEV_CONFLICTS.csv`). Lacunas: `06_tables/NUTEV_LACUNAS_DE_EVIDENCIA.xlsx`.

## 8. Como citar / inserir
- **Referências:** importe `referencias.ris` (Zotero/Mendeley/EndNote) ou use `referencias.bib` (LaTeX).
- **Tabelas:** `NUTEV_DISSERTACAO.xlsx` (abas Resumo, Estudos_incluidos, Por_tema, Por_pais, Por_ano, Por_periodico, Claims_apoiados).
- **Métodos:** texto-base em `08_docs/NUTEV_METHODS_MASTER.md`.
""")

    # --- Word version (best-effort) ---
    if _write_docx(
        out_dir / "RELATORIO_DISSERTACAO.docx",
        title="Revisão de evidências NutEV",
        caveat="Material de apoio — claims e recomendações são candidatos e exigem revisão humana.",
        sections=[("1. Métodos", methods_txt), ("2. Resultados", results_txt), ("3. Avaliação dos artigos", eval_txt)],
        studies=studies,
    ):
        written.append("09_report/RELATORIO_DISSERTACAO.docx")

    # --- index: where each artifact goes in the dissertation ---
    docx_line = "- `RELATORIO_DISSERTACAO.docx` — versão Word (se gerada)\n" if (out_dir / "RELATORIO_DISSERTACAO.docx").exists() else ""
    _write("INDICE.md", f"""# Índice — material para dissertação e artigos

_Gerado em {gen_at}. Pasta: `09_report/`._

## Comece por aqui
1. **`RESUMO_EXECUTIVO.md`** — números-chave (1 página).
2. **`RELATORIO_DISSERTACAO.md`** — relatório completo em capítulos.
{docx_line}3. **`NUTEV_DISSERTACAO.xlsx`** — todas as tabelas em uma planilha (abas).

## Onde usar cada parte
| Seção da dissertação | Use |
|---|---|
| **Métodos** | `08_docs/NUTEV_METHODS_MASTER.md` + `06_tables/NUTEV_PRISMA2020_FLOW.csv` + RELATORIO §1 |
| **Resultados** | `NUTEV_DISSERTACAO.xlsx` (abas) + RELATORIO §2–§4 |
| **Discussão** | RELATORIO §3 (avaliação) e §7 (lacunas/conflitos) — você redige a interpretação |
| **Principais achados** | `achados_principais.csv` (cada achado já vem com artigo + **chave de citação** + DOI) |
| **Referências** | `referencias.ris` (Zotero/Mendeley) ou `referencias.bib` (LaTeX) |
| **Apêndices** | `estudos_incluidos.csv` (todos os {len(corpus)} artigos) + `08_docs/NUTEV_SEARCH_STRATEGY_APPENDIX.md` |

## Como citar (rápido)
1. Importe `referencias.ris` no Zotero/Mendeley (ou `referencias.bib` no Overleaf).
2. Em `achados_principais.csv` / `estudos_incluidos.csv`, cada linha tem **`chave_citacao`** (igual à do `.bib`) e o **DOI**.
3. No seu texto: escreva o achado e cite pela chave/DOI — a referência já está no seu gerenciador.

## Conteúdo
- `achados_principais.csv` — **todos** os achados apoiados, ligados ao artigo + chave de citação
- `estudos_incluidos.csv` — **todos os {len(corpus)} artigos** + avaliação por artigo (com `chave_citacao`)
- `referencias.bib` / `referencias.ris` — {len(corpus)} referências
- `NUTEV_DISSERTACAO.xlsx` — Resumo, Estudos_incluidos, Principais_achados, Por_tema, Por_pais, Por_ano, Por_periodico, Por_workstream

> ⚠️ Tudo é apoio: a validação científica final (claims, recomendações) é sua.
""")

    return {
        "available": True,
        "written": written,
        "n_documents": len(corpus),
        "n_references": len(corpus),
        "report_dir": "09_report",
    }
